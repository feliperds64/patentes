from flask import render_template, request
from estevam_flask import app, db
from openpyxl import load_workbook
from estevam_flask.models import Dados
from datetime import datetime
import pandas as pd
from docx import Document

@app.route('/consulta')
def consulta():
    tabela = Dados.query.all()
    return render_template('consulta.html', tabela=tabela)

@app.route('/', methods=['GET', 'POST'])
def upload_file():
    def substituir_texto(local, cod_processo, nom_marca, nom_titular, cod_especificacao, cod_classe, nome_arquivo, data_referencia):
        document = Document(local + nome_arquivo)

        Dictionary = {'cod_processo': str(cod_processo),
                      'nom_marca': nom_marca,
                      'nom_titular': nom_titular,
                      'cod_especificacao': cod_especificacao,
                      'cod_classe': str(cod_classe)}

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for item in Dictionary:
                            if item in paragraph.text:
                                paragraph.text = paragraph.text.replace(item, Dictionary[item])

        document.save(f'estevam_flask\static\output\{nome_arquivo[:-6]}{nom_marca} - {datetime.strftime(data_referencia, "%d%m%Y")}).docx')


    if request.method == 'POST':
        arquivo = request.files['arquivo']
        workbook = load_workbook(arquivo)
        sheet = workbook.active

        nome_arquivo = arquivo.filename
        nome_arquivo = nome_arquivo[:-5]
        nome_arquivo = nome_arquivo[-10:]
        data_referencia = datetime.strptime(nome_arquivo, '%d.%m.%Y')

        planilha = pd.read_excel(arquivo)

        for row in sheet.iter_rows(min_row=2, values_only=True):
            prop_ter, cod_processo, nom_marca, nom_titular, desc_desp, classe, especificacao = row[0], row[1], row[2], row[3], row[4], row[5], row[6]
            dados = Dados(prop_ter=prop_ter, cod_processo=cod_processo, nom_marca=nom_marca, nom_titular=nom_titular, desc_desp=desc_desp, classe=classe, especificacao=especificacao, data_referencia=data_referencia)
            db.session.add(dados)

        db.session.commit()

        for linha in planilha.index:
            cod_processo = planilha.loc[linha, 'Nº do Processo']
            nom_marca = planilha.loc[linha, 'Nome da Marca']
            nom_titular = planilha.loc[linha, 'Titular']
            cod_especificacao = planilha.loc[linha, 'Especificação']
            cod_classe = planilha.loc[linha, 'Classe']

            local = 'estevam_flask/static/templates/'

            if planilha.loc[linha, 'Próprio/ Terceiro'] == 'P' and planilha.loc[
                linha, 'Descrição do Despacho'] == 'Concessão de registro':
                nome_arquivo = 'Concessão - Caso Próprio ().docx'
            elif planilha.loc[linha, 'Próprio/ Terceiro'] == 'P' and planilha.loc[
                linha, 'Descrição do Despacho'] == 'Deferimento de pedido':
                nome_arquivo = 'Deferimento - Caso Próprio ().docx'
            elif planilha.loc[linha, 'Próprio/ Terceiro'] == '3º' and planilha.loc[
                linha, 'Descrição do Despacho'] == 'Concessão de registro':
                nome_arquivo = 'Deferimento - Caso de Terceiros ().docx'
            elif planilha.loc[linha, 'Próprio/ Terceiro'] == '3º' and planilha.loc[
                linha, 'Descrição do Despacho'] == 'Deferimento de pedido':
                nome_arquivo = 'Concessão - Caso de Terceiros ().docx'

            substituir_texto(data_referencia = data_referencia, local = local, nome_arquivo=nome_arquivo, cod_processo=cod_processo, nom_marca=nom_marca,
                             nom_titular=nom_titular, cod_especificacao=cod_especificacao, cod_classe=cod_classe)

        return 'Arquivo importado com sucesso!'


    return render_template('home.html')